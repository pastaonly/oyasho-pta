from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set default font to Japanese-friendly
style = doc.styles['Normal']
style.font.name = 'Yu Gothic'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

def set_cell_font(cell, size=10.5, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Yu Gothic'
            run.font.size = Pt(size)
            run.bold = bold
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def add_para(text, size=10.5, bold=False, align=None, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    p.paragraph_format.space_after = Pt(space_after)
    return p

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# ===== Header =====
add_para('令和8年4月◯◯日', align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('PTA会員 各位')
add_para('', space_after=0)
add_para('世田谷区立尾山台小学校', align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('校　長　小田正弥', align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('PTA会長　筒井太一朗', align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para('')

# Title
add_para('令和8年度　尾山台小学校PTA臨時総会(書面決議)の結果について',
         size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para('春暖の候、会員の皆様には、益々ご清祥のこととお喜び申し上げます。この度は、PTA臨時総会の書面決議にご回答いただき、誠にありがとうございました。集計結果を下記のとおりご報告いたします。')

add_para('記', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

# ===== 1. 総会の成立 =====
add_para('１．総会の成立', size=12, bold=True, space_after=4)
add_para('PTA会員数 ◯◯◯名（保護者：◯◯◯名　教員：◯◯名）')
add_para('　・議決権行使提出　　17名')
add_para('　・委任提出　　　　103名')
add_para('　・白票提出　　　　　2名')
add_para('')
add_para('PTA規則第6章 総会 第15条(１)により、全会員の5分の1以上（定足数◯◯名）の書面表決・委任をもって、本総会は成立いたしました。（提出総数122名）')
add_para('')

# ===== 2. 議案 =====
add_para('２．議案', size=12, bold=True, space_after=4)
add_para('PTA規則第6章 総会 第15条(２)により、議決は過半数の承認をもって、議案が可決されました。')

table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '議案'
hdr[1].text = '内容'
hdr[2].text = '結果'
hdr[3].text = '内訳'
for c in hdr:
    shade_cell(c, 'D9D9D9')
    set_cell_font(c, bold=True)

row = table.rows[1].cells
row[0].text = '第１号議案'
row[1].text = '規約改正案の承認（令和8年5月改正／任意加入の徹底、会費廃止と学校運営必要経費の新設、学年別活動委員会・学校コミュニティプロジェクトの新設、3口座制への移行 等）'
row[2].text = '可決'
row[3].text = '承認：120　非承認：0\n白票：2　無効(未記入)：0'
for c in row:
    set_cell_font(c)

# Column widths
widths = [Cm(2.3), Cm(9.5), Cm(1.8), Cm(3.5)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

add_para('')
add_para('*注: 承認数120名は、議決権行使17名中の賛成17名と委任状103名を合計したものです。白票2名は、選択肢「白票」を選択された方1名と、氏名記載不備により白票扱いとなった方1名の合計です。', size=9)
add_para('')

# ===== 3. 結果報告サマリー =====
add_para('３．結果報告サマリー', size=12, bold=True, space_after=4)

p = doc.add_paragraph()
run = p.add_run('今回の書面決議は、')
run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
run = p.add_run('反対票0・賛成率100%')
run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5); run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
run = p.add_run('で可決されました。多くの会員の皆様から、改革の方向性に対する賛同と感謝のお言葉を頂戴いたしました。')
run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

add_para('')
add_para('寄せられた主なご意見として、以下の3点が挙げられます。')
for item in [
    '改革の方向性への賛同と、PTA役員への感謝',
    '新制度における役員・委員の選出プロセス（卒対等）についての質問',
    '会計の透明性・予算運用への期待',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

add_para('')
add_para('これらのご意見・ご質問に対し、PTAとして以下の対応に取り組んでまいります。')
for item in [
    '選出プロセスの具体化：学年別活動委員会運営細則に基づき、募集要項・選出方法を年度ごとに明確化し、事前に会員へ通知します',
    '会計の透明性：3口座（学校運営必要経費・学校コミュニティ基金・学校未来基金）ごとの収支を情報公開サイトでリアルタイムに公開し、定期総会で報告します',
    '会員からの問い合わせ窓口：規約・運用に関するご質問は随時お問い合わせフォームで受け付け、個別に回答します',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

add_para('')
add_para('ご意見の詳細と、それに対するPTAの回答は、次章（４．主な意見とその回答）をご覧ください。')
add_para('')

# ===== 4. 主な意見とその回答 =====
add_para('４．主な意見とその回答', size=12, bold=True, space_after=4)

# (1)
add_para('(１) 第１号議案（規約改正案）に関する賛成意見', size=11, bold=True, space_after=4)

opinions_1 = [
    ('賛成意見',
     '・強制ではなく自主的な組織になるということに賛成です。',
     'ご賛同ありがとうございます。今回の改革の中心的な考え方は、まさに「自動加入から任意加入へ」「やらされる活動からやってみたい活動へ」という自主性の尊重です。新制度のもとで、会員一人ひとりが納得して参加できる組織を目指してまいります。'),
    ('賛成意見',
     '・業務が簡素化され、3年生保護者がメインの役員の役割も明確になっているので、今後のスムーズな役員決め、運営に繋がると思います。',
     '前向きなご評価をありがとうございます。学年別活動委員会（旗振り・執行・卒対）への再編により、担当学年・役割・時期が明確になりました。運用を通じて継続的に改善を重ね、持続可能な運営体制を定着させてまいります。'),
    ('賛成意見',
     '・みらい基金にも期待。これまでお支払いいただいたOBにも何かしら伝わるといいなと。',
     '学校未来基金へのご期待をありがとうございます。基金の使途・運用状況は、情報公開サイトおよび定期総会で継続的に報告してまいります。過去にPTA会費をご負担いただいた保護者・OBの皆様にも成果が見える形となるよう、透明性の確保に努めます。'),
]

t1 = doc.add_table(rows=1+len(opinions_1), cols=3)
t1.style = 'Table Grid'
hdr = t1.rows[0].cells
hdr[0].text = '区分'
hdr[1].text = '意見の内容'
hdr[2].text = 'PTAとしての回答'
for c in hdr:
    shade_cell(c, 'D9D9D9')
    set_cell_font(c, bold=True)

for i, (kubun, iken, kotae) in enumerate(opinions_1):
    row = t1.rows[i+1].cells
    row[0].text = kubun
    row[1].text = iken
    row[2].text = kotae
    for c in row:
        set_cell_font(c)

widths1 = [Cm(2.2), Cm(5.8), Cm(9.1)]
for row in t1.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths1[i]

add_para('')

# (2)
add_para('(２) ご意見・ご質問フォームより頂いた主な質問', size=11, bold=True, space_after=4)

opinions_2 = [
    ('役員体制に関する質問',
     '・PTA会長がいなくなるということでしょうか？PTA役員は存続するのであれば、PTA役員の役割はどんなことになりますか？',
     'PTA会長・役員は引き続き存続いたします。改革によって変わるのは「会員構成」と「活動体制」であり、役員会（会長・副会長・会計・会計監査）は規約第14〜15条に基づき、総会議案の作成・予算執行・基金の決裁・学校との委託契約の管理等を担います。詳細は改正後の規約およびポータルサイトをご参照ください。'),
    ('卒対の進め方に関する質問',
     '・6年時の卒対はどのように決めるのか？',
     '卒業対策委員会は、学年別活動委員会運営細則に基づき、6年生の保護者の中から選出します。毎年3月末までに役員会が募集要項（活動内容・所要時間・選出方法）を該当学年の保護者に通知し、立候補を優先、不足する場合は話し合い・互選・くじ引き等の方法で選出します。詳細は細則第3条をご参照ください。'),
    ('予算の透明性に関する期待',
     '・今回アップデートされることで、予算（同窓会費含め）がよりクリアになることを望みます。',
     'ご指摘の通り、予算の透明性は新制度の中核に位置付けております。3口座（学校運営必要経費・学校コミュニティ基金・学校未来基金）の収支は情報公開サイトでリアルタイムに公開し、年度末には繰越状況・将来推計を定期総会で報告します。同窓会費等、関連団体との資金関係についても、必要な範囲で会員の皆様にご説明してまいります。'),
    ('クラスLINE等の連絡手段に関する懸念',
     '・クラスLINEに関して、不安に思っていましたので、何かしら対応があると記載があり、安堵しています。',
     '2026年4月の新年度スタートにあたり、以下の対応を実施いたしました。\n・新1年生：学校支援コーディネーターがクラスごとのグループLINEを作成し、QRコードとリンクを「すぐーる」で配布しました。本活動は2027年度以降も継続してまいります。\n・クラス替えがある3年生・5年生：「すぐーる」にてクラスLINEの存在確認を行いました。既に構築済みであったため、学校支援コーディネーターによるサポートは実施しておりません。\nなお、PTAが管理する個人情報は最小限にとどめ、連絡網の電子化・名簿配布の廃止といった方針を継続します。クラスLINE等の日々の運用は各学年・学級の判断に委ねつつ、PTAとしても必要な情報発信を引き続き行ってまいります。'),
]

t2 = doc.add_table(rows=1+len(opinions_2), cols=3)
t2.style = 'Table Grid'
hdr = t2.rows[0].cells
hdr[0].text = '区分'
hdr[1].text = '意見の内容'
hdr[2].text = 'PTAとしての回答'
for c in hdr:
    shade_cell(c, 'D9D9D9')
    set_cell_font(c, bold=True)

for i, (kubun, iken, kotae) in enumerate(opinions_2):
    row = t2.rows[i+1].cells
    row[0].text = kubun
    row[1].text = iken
    row[2].text = kotae
    for c in row:
        set_cell_font(c)

widths2 = [Cm(3.0), Cm(5.5), Cm(8.6)]
for row in t2.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths2[i]

add_para('')

# (3)
add_para('(３) その他寄せられた感謝・応援のお言葉', size=11, bold=True, space_after=4)
add_para('今回の書面決議では、多くの会員の皆様から改革への賛同と、PTA役員への温かい応援・感謝のお言葉をいただきました。一部を抜粋してご紹介いたします。')

for item in [
    '「スマートな改革をありがとうございました。」',
    '「非常に骨の折れる大規模な改革への着手、心よりお礼申し上げます。旧来と比べ無駄なく、とてもシンプルで分かりやすい組織構成になったと思います。」',
    '「時代にあわせて、今の形に合った内容にアップデートしていただきとても嬉しいです。」',
    '「他の公立小中学校では役員や委員のなりてがおらず、とても苦労していると聞きます。このような改革が広がることを期待しています。」',
    '「アップデートに踏み切ったことがまずとても大きなことだと思いますし、感動しました。子供たち、学校、保護者など多方面のことを考えてたくさんのことを検討下さったことと思います。ありがとうございます。」',
    '「微力ですが、何かお手伝いできる事ありましたらお声掛け下さい。」',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Yu Gothic'; run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

add_para('')
add_para('PTA役員一同、皆様からのお言葉を励みとし、新制度のもとでも子どもたちと学校を支える活動を着実に進めてまいります。今後とも尾山台小学校PTAへのご理解・ご協力を賜りますよう、よろしくお願い申し上げます。')
add_para('')
add_para('以上', align=WD_ALIGN_PARAGRAPH.RIGHT)

output = '/Users/fujiwara/claude/projects/PTA-update/決議/令和8年度　尾山台小学校PTA臨時総会(書面決議)の結果.docx'
doc.save(output)
print(f'Saved: {output}')
