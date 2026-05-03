"""Markdown → Word変換（コピペ用に整形）"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, size=10.5, bold=False):
    run.font.name = 'Yu Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')


def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_para(doc, text, size=10.5, bold=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        # Handle inline bold **text**
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_run_font(run, size=size, bold=True)
            elif part:
                run = p.add_run(part)
                set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part:
            run = p.add_run(part)
            set_run_font(run, size=size, bold=False)


def add_table_from_md(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h.strip())
        set_run_font(run, bold=True)
        shade_cell(hdr[i], 'D9D9D9')
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            text = val.strip()
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run_font(run)


def parse_table(lines, idx):
    """Parse markdown table starting at idx. Return (header, rows, next_idx)."""
    header_line = lines[idx]
    header = [c.strip() for c in header_line.strip().strip('|').split('|')]
    idx += 2  # skip separator
    rows = []
    while idx < len(lines) and lines[idx].lstrip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        rows.append(row)
        idx += 1
    return header, rows, idx


def md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Headers
        if stripped.startswith('#### '):
            add_para(doc, stripped[5:], size=11, bold=True, space_after=4)
        elif stripped.startswith('### '):
            add_para(doc, stripped[4:], size=12, bold=True, space_after=4)
        elif stripped.startswith('## '):
            add_para(doc, stripped[3:], size=13, bold=True, space_after=6)
        elif stripped.startswith('# '):
            add_para(doc, stripped[2:], size=16, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
        # Tables
        elif stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
            header, rows, i = parse_table(lines, i)
            add_table_from_md(doc, header, rows)
            add_para(doc, '', space_after=2)
            continue
        # Bullets
        elif stripped.startswith('- '):
            add_bullet(doc, stripped[2:])
        # Bold-prefixed line, e.g. "**foo:** bar"
        else:
            add_para(doc, stripped)

        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    base = Path('/Users/fujiwara/claude/projects/PTA-update/3年生執行委員-運営ルール')
    for md in sorted(base.glob('*.md')):
        out = md.with_suffix('.docx')
        md_to_docx(md, out)
